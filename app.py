import os
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re

app = Flask(__name__)
CORS(app)

# مائیکروسافٹ ورڈ کو آفیشل اردو/عربی اسکرپٹ (RTL) پر لاک کرنے کا فنکشن
def apply_strict_word_rtl(paragraph, font_name, font_size_pt, lang_code):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    if paragraph.runs:
        for run in paragraph.runs:
            rPr = run._r.get_or_add_rPr()
            
            rtl_element = OxmlElement('w:rtl')
            rtl_element.set(qn('w:val'), '1')
            rPr.append(rtl_element)
            
            lang = OxmlElement('w:lang')
            lang.set(qn('w:bidi'), lang_code)
            rPr.append(lang)
            
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:cs'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rPr.append(rFonts)
            
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size_pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size_pt * 2)))
            rPr.append(szCs)

# قرآنی اعراب اور علامات چیک کرنے کا فنکشن
def is_arabic_line(text):
    arabic_char_count = len(re.findall(r'[\u064b-\u065f\u0671\u06d6-\u06dc]', text))
    if arabic_char_count > 0:
        return True
    return False

# 🔥 ورڈ 10 کی پی ڈی ایف خامیوں کو جڑ سے ختم کرنے کا جادوئی انجن
def fix_word10_urdu_layout(text):
    lines = text.split('\n')
    fixed_lines = []
    
    for line in lines:
        if not line.strip():
            fixed_lines.append("")
            continue
            
        # 1. اصلی الفاظ بلاکس کو تلاش کرنا (جو ڈبل یا اس سے زیادہ اسپیس سے الگ ہیں)
        word_blocks = re.split(r'\s{2,}', line.strip())
        fixed_blocks = []
        
        for block in word_blocks:
            # حروف کے درمیان موجود فیک سنگل اسپیس کو ختم کرنا
            fused_word = block.replace(" ", "")
            if not fused_word:
                continue
                
            # اگر بلاک صرف نمبر (جیسے 2020) یا انگلش ہے تو اسے سیدھا رکھنا ہے
            if re.match(r'^[A-Za-z0-9\W]+$', fused_word) and not re.search(r'[\u0600-\u06FF]', fused_word):
                fixed_blocks.append(fused_word)
            else:
                # اگر اردو/عربی ہے تو ورڈ 10 کی الٹی کوڈنگ کو سیدھا کرنے کے لیے حروف کو ریورس کرنا
                fixed_blocks.append(fused_word[::-1])
                
        # 2. چونکہ پورا جملہ بائیں سے دائیں الٹا تھا، اس لیے الفاظ کی ترتیب کو ریورس کرنا
        fixed_blocks.reverse()
        
        # 3. صاف ستھرے الفاظ کو سنگل اسپیس کے ساتھ جوڑنا
        fixed_lines.append(" ".join(fixed_blocks))
        
    return "\n".join(fixed_lines)

@app.route('/')
def home():
    return "Perfect Word 10 Urdu Logical Converter Backend is Live!"

@app.route('/convert', methods=['POST'])
def convert_pdf_to_word():
    if 'pdfFile' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    file = request.files['pdfFile']
    if file.filename == '':
        return jsonify({"error": "Empty file name"}), 400

    try:
        pdf_bytes = file.read()
        extracted_text = ""
        
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"

        if not extracted_text.strip():
            return jsonify({"error": "Is PDF me koi text nahi mila."}), 400

        # 🔥 مائیکروسافٹ ورڈ 10 کے ٹیکسٹ کو لائیو سیدھا اور فلٹر کرنا
        perfect_logical_text = fix_word10_urdu_layout(extracted_text)

        doc = Document()
        
        lines = perfect_logical_text.split('\n')
        for line in lines:
            if line.strip():
                p = doc.add_paragraph()
                p.add_run(line) # اب خالص، جڑا ہوا اور اصلی لاجیکل ٹیکسٹ ورڈ میں جائے گا
                
                if is_arabic_line(line):
                    font_name = 'Traditional Arabic'
                    font_size = 16
                    lang_code = 'ar-SA'
                else:
                    font_name = 'Noto Nastaliq Urdu'
                    font_size = 14
                    lang_code = 'ur-PK'
                
                apply_strict_word_rtl(p, font_name, font_size, lang_code)

        word_io = io.BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        
        return send_file(
            word_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="Urdu_Perfect_Logical_Fixed.docx"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
